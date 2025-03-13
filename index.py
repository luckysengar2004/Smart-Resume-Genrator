import streamlit as st
import openai
import os
import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
import google.generativeai as genai

# Loading environment variables
load_dotenv()
GOOGLE_API_KEY = os.getenv("API  -- KEY") # your api key

genai.configure(api_key=GOOGLE_API_KEY)

# Streamlit UI Setup
st.set_page_config(page_title="SmartResume Generator", page_icon="📄", layout="wide")
st.title("📄 SmartResume Generator")
st.markdown("Create a professional resume tailored for every opportunity.")

# Sidebar for settings
with st.sidebar:
    st.header("Settings")
    theme = st.selectbox("Select Theme", ["Classic", "Modern", "Minimalist"])
    font_size = st.slider("Font Size", 10, 20, 12)
    st.markdown("---")
    st.write("🔹 **Ensure all fields are filled for best results.**")

# Tabs for a structured layout
tabs = st.tabs(["Personal Info", "Experience", "Skills & Education", "Review & Generate"])

# Personal Information
tabs[0].subheader("Personal Information")
name = tabs[0].text_input("Full Name")
email = tabs[0].text_input("Email")
phone = tabs[0].text_input("Phone")
linkedin = tabs[0].text_input("LinkedIn URL")
summary = tabs[0].text_area("Professional Summary")

# Experience
tabs[1].subheader("Work Experience")
experience = []
num_experiences = tabs[1].number_input("Number of experiences", min_value=1, max_value=10, step=1)

for i in range(num_experiences):
    with tabs[1].expander(f"Experience {i+1}"):
        job_title = st.text_input(f"Job Title {i+1}", key=f"job_title_{i}")
        company = st.text_input(f"Company {i+1}", key=f"company_{i}")
        duration = st.text_input(f"Duration {i+1}", key=f"duration_{i}")
        description = st.text_area(f"Description {i+1}", key=f"description_{i}")
        experience.append({
            "job_title": job_title,
            "company": company,
            "duration": duration,
            "description": description
        })

# Skills & Education
tabs[2].subheader("Skills & Education")
skills = tabs[2].text_area("Skills (comma-separated)")

degree = tabs[2].text_input("Degree")
university = tabs[2].text_input("University")
grad_year = tabs[2].text_input("Graduation Year")

certifications = tabs[2].text_area("Certifications")
languages = tabs[2].text_area("Languages Spoken")
projects = tabs[2].text_area("Key Projects")

# Resume Generation Function
def generate_resume():
    try:
        model = genai.GenerativeModel("gemini-pro")
        prompt = f"""
        Generate a professional resume with the following details:
        Name: {name}
        Email: {email}
        Phone: {phone}
        LinkedIn: {linkedin}
        Summary: {summary}
        Experience: {experience}
        Skills: {skills}
        Education: {degree}, {university}, {grad_year}
        Certifications: {certifications}
        Languages Spoken: {languages}
        Key Projects: {projects}
        Format professionally.
        """
        
        response = model.generate_content(prompt)
        return response.text if response else "Error generating resume."
    except Exception as e:
        return f"An error occurred: {str(e)}"

# Save Resume to Word Document
def save_to_word(resume_text):
    try:
        doc = docx.Document()
        doc.add_heading(name, level=1)

        contact_info = f"{email} | {phone} | {linkedin}"
        doc.add_paragraph(contact_info).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_heading("Professional Summary", level=2)
        doc.add_paragraph(summary)

        doc.add_heading("Work Experience", level=2)
        for exp in experience:
            doc.add_paragraph(f"{exp['job_title']} at {exp['company']} ({exp['duration']})", style="Heading 3")
            doc.add_paragraph(exp['description'])

        doc.add_heading("Skills", level=2)
        doc.add_paragraph(skills)

        doc.add_heading("Education", level=2)
        doc.add_paragraph(f"{degree}, {university} ({grad_year})")

        if certifications:
            doc.add_heading("Certifications", level=2)
            doc.add_paragraph(certifications)
        if languages:
            doc.add_heading("Languages Spoken", level=2)
            doc.add_paragraph(languages)
        if projects:
            doc.add_heading("Key Projects", level=2)
            doc.add_paragraph(projects)

        file_path = "Generated_Resume.docx"
        doc.save(file_path)
        return file_path
    except Exception as e:
        return f"Error saving resume: {str(e)}"

# Review & Generate
tabs[3].subheader("Review & Generate Resume")
if tabs[3].button("Generate Resume ✨"):
    resume_text = generate_resume()
    tabs[3].text_area("Generated Resume Preview", resume_text, height=400)
    file_path = save_to_word(resume_text)
    if file_path:
        with open(file_path, "rb") as f:
            tabs[3].download_button("📥 Download Resume", f, file_name="Generated_Resume.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
